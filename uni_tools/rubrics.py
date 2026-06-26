from shutil import copy2
from docx import Document
from docx.shared import Pt, Inches
from .signatures import get_student_signature
import os


def fill_rubric(template_path: str, output_path: str, student_name: str, student_mat: str, fecha: str = "30/06/2026"):
    copy2(template_path, output_path)
    doc = Document(output_path)

    mt = doc.tables[0]
    n1 = mt.rows[1].cells[0].tables[1]
    tc = n1.rows[5].cells[0]
    nt = tc.tables[0]

    name_p = nt.rows[0].cells[1].paragraphs[0]
    name_p.clear()
    r = name_p.add_run(student_name)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.bold = True

    date_p = nt.rows[0].cells[1].add_paragraph()
    date_p.clear()
    rf = date_p.add_run(fecha)
    rf.font.name = "Times New Roman"
    rf.font.size = Pt(10)

    sig_cell = nt.rows[1].cells[1]
    sig_cell.paragraphs[0].clear()

    sig_bytes = get_student_signature(student_name)
    if sig_bytes:
        media_dir = os.path.join(os.path.dirname(output_path), "media")
        os.makedirs(media_dir, exist_ok=True)
        sig_file = os.path.join(media_dir, f"sig_{student_mat}.png")
        with open(sig_file, "wb") as f:
            f.write(sig_bytes)
        sig_cell.paragraphs[0].add_run().add_picture(sig_file, width=Inches(1.5))

    doc.save(output_path)
    return output_path
