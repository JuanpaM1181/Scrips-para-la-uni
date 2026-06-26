from shutil import copy2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .formatting import set_paragraph_font


def copy_template(dst: str) -> str:
    copy2(TEMPLATE_COVER, dst)
    return dst


def fill_cover(
    doc: Document,
    students: list,
    materia: str,
    periodo: str,
    unidad: str,
    titulo: str,
    tipo_evidencia: str,
    fecha: str,
    carrera: str = "Ingenieria de Software.",
):
    paragraph_map = {
        5: "Nombre del alumno – Matricula",
        6: f"{students[0]['name']} – {students[0]['matricula']}",
        7: f"{students[1]['name']} - {students[1]['matricula']}" if len(students) > 1 else "",
        9: "Materia.",
        10: materia,
        12: "Cuatrimestre / periodo escolar",
        13: periodo,
        15: "Unidad",
        16: unidad,
        17: "Nombre de la practica / proyecto.",
        18: titulo,
        20: "Tipo de evidencia",
        21: tipo_evidencia,
        23: "Plan de estudios.",
        24: carrera,
        28: fecha,
    }

    for i, p in enumerate(doc.paragraphs):
        if i in paragraph_map and paragraph_map[i] is not None:
            runs = p.runs
            if runs:
                runs[0].text = paragraph_map[i]
                for extra in runs[1:]:
                    extra.text = ""
        set_paragraph_font(p, "Times New Roman", Pt(12))
        if i < 29:
            p.paragraph_format.line_spacing = 1.0
            try:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    last_cover = doc.paragraphs[28]
    run_br = last_cover.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_br._r.append(br)


from .config import TEMPLATE_COVER
