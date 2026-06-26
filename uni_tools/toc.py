from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def insert_toc(doc):
    from .formatting import add_heading_custom

    add_heading_custom(doc, "Indice", 1)
    toc_para = doc.add_paragraph()
    toc_para.paragraph_format.line_spacing = 1.5
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:r")
    placeholder_pt = OxmlElement("w:t")
    placeholder_pt.set(qn("xml:space"), "preserve")
    placeholder_pt.text = "[Actualice la tabla de contenidos en Word]"
    placeholder.append(placeholder_pt)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_toc = toc_para.add_run()
    run_toc._r.append(fld_char_begin)
    run_toc._r.append(instr_text)
    run_toc._r.append(fld_char_separate)
    run_toc._r.append(placeholder)
    run_toc._r.append(fld_char_end)
    run_toc.font.name = "Times New Roman"
    run_toc.font.size = Pt(12)
    return toc_para
