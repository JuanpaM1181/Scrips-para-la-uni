from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_paragraph_font(paragraph, font_name="Times New Roman", font_size=Pt(12)):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = font_size


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if p.runs:
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(14) if level == 1 else Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(1.27)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    prefix = "- " if level == 0 else "  - "
    run = p.add_run(prefix + text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p
