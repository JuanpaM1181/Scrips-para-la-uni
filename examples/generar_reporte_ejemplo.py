from uni_tools import add_heading_custom, add_body, add_bullet, insert_toc, convert_to_pdf
from uni_tools.cover import copy_template, fill_cover
from uni_tools.config import STUDENTS
from docx import Document

dst = "/tmp/Reporte_Ejemplo.docx"
copy_template(dst)
doc = Document(dst)

fill_cover(
    doc=doc,
    students=STUDENTS,
    materia="Programacion para Moviles II",
    periodo="9o (Mayo - Agosto 2026)",
    unidad="2",
    titulo="Evidencia de Desempeno 1 - CRUD con Base de Datos Web",
    tipo_evidencia="Reporte Tecnico de Arquitectura + Aplicacion Funcional",
    fecha="Tapachula, Chiapas; a 30 de junio del 2026",
)

insert_toc(doc)

add_heading_custom(doc, "Introduccion", 1)
add_body(doc, "Texto de la introduccion aqui...")

add_heading_custom(doc, "Seccion 1. Marco Teorico", 1)
add_body(doc, "Contenido de la seccion...")
add_bullet(doc, "Punto importante 1")
add_bullet(doc, "Punto importante 2")

add_heading_custom(doc, "Conclusion", 1)
add_body(doc, "Conclusion del trabajo...")

doc.save(dst)
convert_to_pdf(dst)
print("Reporte generado exitosamente.")
